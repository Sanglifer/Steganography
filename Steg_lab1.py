import textwrap
from docx import Document
from docx_utils.flatten import opc_to_flat_opc
import xml.etree.ElementTree as ET


def KOI8R_WINDOWS1251_CP866(bits, encoding, errors='surrogatepass'):
    while len(bits) % 8 != 0:
        bits = bits[:-1]
    n = int(bits, 2)
    return n.to_bytes((n.bit_length() + 7) // 8, 'big').decode(encoding, errors) or '\0'


def Baudot_MTK2(code):
    Switch = {
        "Let": "00000",
        "Dig": "11011",
        "SP": "00100"
    }

    Letters = {
        "А": "00011",
        "Б": "11001",
        "Ц": "01110",
        "Д": "01001",
        "Е": "00001",
        "Ф": "01101",
        "Г": "11010",
        "Х": "10100",
        "И": "00110",
        "Й": "01011",
        "К": "01111",
        "Л": "10010",
        "М": "11100",
        "Н": "01100",
        "О": "11000",
        "П": "10110",
        "Я": "10111",
        "Р": "01010",
        "С": "00101",
        "Т": "10000",
        "У": "00111",
        "Ж": "11110",
        "В": "10011",
        "Ь": "11101",
        "Ы": "10101",
        "З": "10001"
    }

    Digits = {
        "-": "00011",
        "?": "11001",
        ":": "01110",
        "Кто там?": "01001",
        "3": "00001",
        "Э": "01101",
        "Ш": "11010",
        "Щ": "10100",
        "8": "00110",
        "Ю": "01011",
        "(": "01111",
        ")": "10010",
        ".": "11100",
        ",": "01100",
        "9": "11000",
        "0": "10110",
        "1": "10111",
        "Ч": "01010",
        "'": "00101",
        "5": "10000",
        "7": "00111",
        "=": "11110",
        "2": "10011",
        "/": "11101",
        "6": "10101",
        "+": "10001"
    }

    def list2string(data):
        return ''.join(map(str, data))

    split_bits = textwrap.wrap(code, 5)

    decoded = []
    Let = False
    Dig = False
    for bits in split_bits:
        if bits == Switch["Let"]:
            Let = True
            Dig = False
        elif bits == Switch["Dig"]:
            Dig = True
            Let = False
        elif bits == Switch["SP"]:
            decoded.extend(" ")
        else:
            if Let:
                if bits in Switch.values():
                    if bits == Switch["Dig"]:
                        Let = False
                        Dig = True
                    else:
                        Dig = False
                        Let = True

                else:
                    for k, v in Letters.items():
                        if v == bits:
                            decoded.extend(k)
            elif Dig:
                if bits in Switch.values():
                    if bits == Switch["Let"]:
                        Dig = False
                        Let = True
                    else:
                        Dig = True
                        Let = False
                else:
                    for k, v in Digits.items():
                        if v == bits:
                            decoded.extend(k)

    decoded = list2string(decoded)
    return decoded


def CheckSteg():
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if str(run.font.color.rgb) != "000000":
                code = TextColor()
                print("Информация скрыта посредством изменения цвета символов")
                return code
            elif str(run.font.size.pt) != "12.0":
                code = TextSize()
                print("Информация скрыта посредством изменения размера шрифта")
                return code
            else:
                break
    for run in root.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"):
        if run.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w") is not None:
            count = ForXML()
            code = XML(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w", count)
            print("Информация скрыта посредством изменения масштаба шрифта")
            return code
        elif run.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd") is not None:
            count = ForXML()
            code = XML(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd", count)
            print("Информация скрыта посредством изменения цвета фона")
            return code
        else:
            count = ForXML()
            code = XML(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing", count)
            print("Информация скрыта посредством изменения межсимвольного интервала")
            return code


def TextColor():
    code = ""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if len(paragraph.text) == len(run.text):
                return code
            else:
                if str(run.font.color.rgb) == "000000":
                    code += ("0" * len(run.text))
                else:
                    code += ("1" * len(run.text))


def TextSize():
    code = ""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if len(paragraph.text) == len(run.text):
                return code
            else:
                if str(run.font.size.pt) == "12.0":
                    code += ("0" * len(run.text))
                else:
                    code += ("1" * len(run.text))


def ForXML():
    count = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if len(paragraph.text) == len(run.text):
                return count
            else:
                count += 1


def XML(attribute, countDOCX):
    countXML = 0
    code = ""
    for run in root.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"):
        countXML += 1
        if countXML == countDOCX:
            return code
        else:
            if run.find(attribute) is not None:
                code += ("1" * len(run.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t").text))
            else:
                code += ("0" * len(run.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t").text))


def Decode(code):
    print("<<Декодирование сообщения>>")
    print("KOI8-R:", KOI8R_WINDOWS1251_CP866(code, "KOI8-R"))
    print("Windows-1251:", KOI8R_WINDOWS1251_CP866(code, "Windows-1251"))
    print("cp866:", KOI8R_WINDOWS1251_CP866(code, "cp866"))
    print("Код Бодо (МТК-2):", Baudot_MTK2(code))


def OpenDOCX_CreateXML():
    doc = Document(docx_file)
    opc_to_flat_opc(docx_file, xml_file)
    root = ET.parse(xml_file).getroot()
    return doc, root


docx_file = "C:\Steg\Variant11_scale.docx"
xml_file = "C:\Steg\Variant11_scale.xml"
doc, root = OpenDOCX_CreateXML()
code = CheckSteg()
print("Закодированное сообщение:", code)
Decode(code)