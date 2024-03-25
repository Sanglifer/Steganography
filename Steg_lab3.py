from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING


def text_to_bits(text, encoding='Windows-1251', errors='surrogatepass'):
    bits = bin(int.from_bytes(text.encode(encoding, errors), 'big'))[2:]
    return bits.zfill(8 * ((len(bits) + 7) // 8))


def text_from_bits(bits, encoding='Windows-1251', errors='surrogatepass'):
    n = int(bits, 2)
    return n.to_bytes((n.bit_length() + 7) // 8, 'big').decode(encoding, errors) or '\0'


def Text_to_List():
    docx_emp = Document(docx_empty)
    count = 0
    for paragraph in docx_emp.paragraphs:
        for run in paragraph.runs:
            count += 1
    return count


def CheckSize():
    print("Скрываемое сообщение:", message)
    if 8*len(message) <= count:
        print("Длина сообщения удовлетворяет размеру контейнера")
        return True
    else:
        print("Длина сообщения не удовлетворяет размеру контейнера")
        return False


def Hide():
    m = ""
    for symb in message:
        m += text_to_bits(symb)
    print("Сообщение в двоичном формате:", m)

    doc = Document(docx_empty)
    docx_new = Document()
    index = 0
    n = False
    for paragraph in doc.paragraphs:
        p = docx_new.add_paragraph()
        for run in paragraph.runs:
            lst = list(run.text)
            if lst == ["\n"]:
                n = True
            else:
                if index < len(m):
                    if n:
                        lst.insert(0, "\n")
                        n = False
                    if m[index] == "0":
                        r = p.add_run("".join(lst) + " ")
                        if index == 0:
                            r.font.name = "Helvetica"
                            r.font.size = Pt(19.5)
                            r.font.color.rgb = RGBColor(51, 51, 51)
                            fmt = p.paragraph_format
                            fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            fmt.space_after = Pt(11.25)
                        else:
                            r.font.name = "Helvetica"
                            r.font.size = Pt(13.5)
                            r.font.color.rgb = RGBColor(51, 51, 51)
                            fmt = p.paragraph_format
                            fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                            fmt.line_spacing = Pt(20.4)
                            fmt.space_after = Pt(18.75)
                    if m[index] == "1":
                        r = p.add_run("".join(lst) + "  ")
                        if index == 0:
                            r.font.name = "Helvetica"
                            r.font.size = Pt(19.5)
                            r.font.color.rgb = RGBColor(51, 51, 51)
                            fmt = p.paragraph_format
                            fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            fmt.space_after = Pt(11.25)
                        else:
                            r.font.name = "Helvetica"
                            r.font.size = Pt(13.5)
                            r.font.color.rgb = RGBColor(51, 51, 51)
                            fmt = p.paragraph_format
                            fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                            fmt.line_spacing = Pt(20.4)
                            fmt.space_after = Pt(18.75)
                    index += 1
                else:
                    if n:
                        lst.insert(0, "\n")
                        n = False
                    r = p.add_run("".join(lst))
                    r.font.name = "Helvetica"
                    r.font.size = Pt(13.5)
                    r.font.color.rgb = RGBColor(51, 51, 51)
                    fmt = p.paragraph_format
                    fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                    fmt.line_spacing = Pt(20.4)
                    fmt.space_after = Pt(18.75)
    docx_new.save(docx_full)
    print("Сообщение сокрыто в файле", docx_full)


def FromFile():
    print("Бит сокрыт в строке (ДА/НЕТ):")
    doc = Document(docx_full)
    Y = []
    N = []
    bits = ""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            lst = list(run.text)
            if '\n' in lst:
                lst.remove('\n')
            if (lst[len(lst)-1] == ' ') and (lst[len(lst)-2] != ' '):
                bits += "0"
                l = "".join(lst)
                Y.append(l)
                print(l, "<--", "ДА")
            elif (lst[len(lst)-1] == ' ') and (lst[len(lst)-2] == ' '):
                bits += "1"
                l = "".join(lst)
                Y.append(l)
                print(l, "<--", "ДА")
            else:
                l = "".join(lst)
                N.append(l)
                print(l, "<--", "НЕТ")

    message_from_file = text_from_bits(bits)
    print("Подмножество Y:", Y)
    print("Подмножество N:", N)
    print("Сообщение, сокрытое в контейнере:", message_from_file)



message = "Секрет"
docx_empty = "C:\Steg\lab_3.docx"
docx_full = "C:\Steg\lab_3_full.docx"
count = Text_to_List()
if CheckSize():
    Hide()
    FromFile()