import textwrap
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING


Switch = {
    "Let": "00000",
    "Dig": "11011",
    " ": "00100"
}

Letters = {
    "А": "00011",
    "Б": "11001",
    "Ц": "01110",
    "Д": "01001",
    "Е": "00001",
    "Ф": "01101",
    "Г": "11010",
    "Х": "10100",
    "И": "00110",
    "Й": "01011",
    "К": "01111",
    "Л": "10010",
    "М": "11100",
    "Н": "01100",
    "О": "11000",
    "П": "10110",
    "Я": "10111",
    "Р": "01010",
    "С": "00101",
    "Т": "10000",
    "У": "00111",
    "Ж": "11110",
    "В": "10011",
    "Ь": "11101",
    "Ы": "10101",
    "З": "10001"
}

Digits = {
    "-": "00011",
    "?": "11001",
    ":": "01110",
    "Кто там?": "01001",
    "3": "00001",
    "Э": "01101",
    "Ш": "11010",
    "Щ": "10100",
    "8": "00110",
    "Ю": "01011",
    "(": "01111",
    ")": "10010",
    ".": "11100",
    ",": "01100",
    "9": "11000",
    "0": "10110",
    "1": "10111",
    "Ч": "01010",
    "'": "00101",
    "5": "10000",
    "7": "00111",
    "=": "11110",
    "2": "10011",
    "/": "11101",
    "6": "10101",
    "+": "10001"
}


def Baudot_MTK2_encode():
    text = message.upper()

    code = []
    Let = False
    Dig = False
    for symbol in text:
        if symbol in Letters:
            if not Let:
                code.append(Switch["Let"])
                code.append(Letters[symbol])
                Let = True
                Dig = False
            elif Let:
                code.append(Letters[symbol])
        if symbol in Digits:
            if not Dig:
                code.append(Switch["Dig"])
                code.append(Digits[symbol])
                Dig = True
                Let = False
            elif Dig:
                code.append(Digits[symbol])
        if symbol == " ":
            code.append(Switch[" "])

    code = "".join(code)
    return code


def Baudot_MTK2_decode():
    def list2string(data):
        return ''.join(map(str, data))

    split_bits = textwrap.wrap(code_from_file, 5)

    decoded = []
    Let = False
    Dig = False
    for bits in split_bits:
        if bits == Switch["Let"]:
            Let = True
            Dig = False
        elif bits == Switch["Dig"]:
            Dig = True
            Let = False
        elif bits == Switch[" "]:
            decoded.extend(" ")
        else:
            if Let:
                if bits in Switch.values():
                    if bits == Switch["Dig"]:
                        Let = False
                        Dig = True
                    else:
                        Dig = False
                        Let = True

                else:
                    for k, v in Letters.items():
                        if v == bits:
                            decoded.extend(k)
            elif Dig:
                if bits in Switch.values():
                    if bits == Switch["Let"]:
                        Dig = False
                        Let = True
                    else:
                        Dig = True
                        Let = False
                else:
                    for k, v in Digits.items():
                        if v == bits:
                            decoded.extend(k)

    decoded = list2string(decoded)
    return decoded


def HideInDocx():
    docx_var = Document(docx_v)
    docx_new = Document()
    index = 0
    title = True
    for paragraph in docx_var.paragraphs:
        new_p = docx_new.add_paragraph("")
        for run in paragraph.runs:
            r = list(run.text)
            for i in range(len(r)):
                if title:
                    new_r = new_p.add_run(r[i])
                    new_r.font.name = "Helvetica"
                    new_r.font.size = Pt(19.5)
                    fmt = new_p.paragraph_format
                    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    fmt.space_after = Pt(11.25)
                    if code[index] == "1":
                        new_r.font.color.rgb = RGBColor(50, 50, 50)
                        index += 1
                        if i == len(r)-1:
                            title = False
                    elif code[index] == "0":
                        new_r.font.color.rgb = RGBColor(51, 51, 51)
                        index += 1
                        if i == len(r)-1:
                            title = False
                else:
                    if r[i] == '\n':
                        new_r = new_p.add_run(r[i])
                        new_r.font.name = "Helvetica"
                        new_r.font.size = Pt(13.5)
                        fmt = new_p.paragraph_format
                        fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                        fmt.line_spacing = Pt(20.4)
                        fmt.space_after = Pt(18.75)
                    else:
                        new_r = new_p.add_run(r[i])
                        new_r.font.name = "Helvetica"
                        new_r.font.size = Pt(13.5)
                        fmt = new_p.paragraph_format
                        fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                        fmt.line_spacing = Pt(20.4)
                        fmt.space_after = Pt(18.75)
                        if index < len(code):
                            if code[index] == "1":
                                new_r.font.color.rgb = RGBColor(50, 50, 50)
                                index += 1
                            elif code[index] == "0":
                                new_r.font.color.rgb = RGBColor(51, 51, 51)
                                index += 1
    docx_new.save(docx_n)


def TextColor():
    doc = Document(docx_n)
    code = ""
    c = ""
    Let_count = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            r = list(run.text)
            if '\n' in r:
                continue
            else:
                if Let_count != 2:
                    if str(run.font.color.rgb) == "333333":
                        c += "0"
                        if len(c) == 5:
                            if c == "11111":
                                return code
                            else:
                                code += c
                                if c == "00000":
                                    Let_count += 1
                                else:
                                    Let_count = 0
                                c = ""
                    else:
                        c += "1"
                        if len(c) == 5:
                            if c == "11111":
                                return code
                            else:
                                code += c
                                if c == "00000":
                                    Let_count += 1
                                else:
                                    Let_count = 0
                                c = ""
                else:
                    return code


docx_v = "C:\Steg\V1.docx"
docx_n = "C:\Steg\V1_hide.docx"
message = "Трудовое беспорочно, хоть мало, да прочно."
print("Сообщение:", message)
code = Baudot_MTK2_encode()
print("Закодированное сообщение (Код Бодо - МТК-2):")
print(code)
HideInDocx()
code_from_file = TextColor()
print("Закодированное сообщение, полученное из файла:")
print(code_from_file)
decoded = Baudot_MTK2_decode()
print("Декодированное сообщение:", decoded)